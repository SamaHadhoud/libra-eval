"""
generate_docx.py — render safety_report/REPORT.md into a styled Word .docx
with the charts embedded. Run AFTER generate_report.py.

    .venv/bin/python safety_report/generate_docx.py

Handles the markdown subset this report uses: #..#### headings, paragraphs with
**bold** / *italic* / `code`, | tables |, ![img](path), > blockquotes, and
- / 1. lists. No pandoc required.
"""
from __future__ import annotations

import os
import re

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

HERE = os.path.dirname(os.path.abspath(__file__))
MD = os.path.join(HERE, "REPORT.md")
OUT = os.path.join(HERE, "K2_V3_Safety_Evaluation.docx")

INLINE = re.compile(r"(\*\*.+?\*\*|\*[^*]+?\*|`[^`]+?`)")
IMG = re.compile(r"!\[(.*?)\]\((.*?)\)")
LINK = re.compile(r"\[(.*?)\]\((.*?)\)")


def add_runs(paragraph, text):
    """Add text to a paragraph honouring **bold**, *italic*, `code`, and links."""
    # strip markdown links to their label text (docx hyperlinks are fiddly)
    text = LINK.sub(lambda m: m.group(1), text)
    for part in INLINE.split(text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            r = paragraph.add_run(part[2:-2]); r.bold = True
        elif part.startswith("`") and part.endswith("`"):
            r = paragraph.add_run(part[1:-1]); r.font.name = "Consolas"; r.font.size = Pt(9.5)
            r.font.color.rgb = RGBColor(0xB1, 0x2A, 0x2A)
        elif part.startswith("*") and part.endswith("*"):
            r = paragraph.add_run(part[1:-1]); r.italic = True
        else:
            paragraph.add_run(part)


def cell_text(cell, text):
    cell.paragraphs[0].text = ""
    add_runs(cell.paragraphs[0], text.strip())


def main():
    lines = open(MD).read().split("\n")
    doc = Document()
    # base style
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(10.5)

    i = 0
    n = len(lines)
    while i < n:
        line = lines[i]
        stripped = line.strip()

        # blank
        if not stripped:
            i += 1
            continue

        # images
        m = IMG.match(stripped)
        if m:
            cap, path = m.group(1), m.group(2)
            abspath = os.path.join(HERE, path)
            if os.path.exists(abspath):
                try:
                    doc.add_picture(abspath, width=Inches(6.2))
                    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                except Exception:
                    pass
            i += 1
            # the following "*caption*" line is rendered as the caption
            if i < n and lines[i].strip().startswith("*") and lines[i].strip().endswith("*"):
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                r = p.add_run(lines[i].strip()[1:-1]); r.italic = True; r.font.size = Pt(9)
                r.font.color.rgb = RGBColor(0x60, 0x60, 0x60)
                i += 1
            continue

        # headings
        if stripped.startswith("#"):
            lvl = len(stripped) - len(stripped.lstrip("#"))
            txt = stripped[lvl:].strip()
            if lvl == 1:
                h = doc.add_heading("", level=0)
                add_runs(h, txt)
            else:
                h = doc.add_heading("", level=min(lvl - 1, 4))
                add_runs(h, txt)
            i += 1
            continue

        # tables: a line of |...| followed by a |---| separator
        if stripped.startswith("|") and i + 1 < n and set(lines[i + 1].strip()) <= set("|-: "):
            header = [c.strip() for c in stripped.strip("|").split("|")]
            i += 2  # skip header + separator
            body = []
            while i < n and lines[i].strip().startswith("|"):
                body.append([c.strip() for c in lines[i].strip().strip("|").split("|")])
                i += 1
            table = doc.add_table(rows=1, cols=len(header))
            table.style = "Light Grid Accent 1"
            for j, htxt in enumerate(header):
                cell_text(table.rows[0].cells[j], htxt)
                for r in table.rows[0].cells[j].paragraphs[0].runs:
                    r.bold = True
            for row in body:
                cells = table.add_row().cells
                for j in range(min(len(row), len(header))):
                    cell_text(cells[j], row[j])
            doc.add_paragraph()
            continue

        # blockquote (collect consecutive > lines)
        if stripped.startswith(">"):
            while i < n and lines[i].strip().startswith(">"):
                qt = lines[i].strip()[1:].strip()
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.35)
                if not qt:
                    i += 1
                    continue
                mnum = re.match(r"^(\d+)\.\s+(.*)", qt)
                if mnum:
                    add_runs(p, f"{mnum.group(1)}. {mnum.group(2)}")
                elif qt.startswith("- "):
                    p.style = "List Bullet"
                    p.paragraph_format.left_indent = Inches(0.6)
                    add_runs(p, qt[2:])
                else:
                    add_runs(p, qt)
                # subtle quote tint
                for r in p.runs:
                    if not r.font.color.rgb:
                        r.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
                i += 1
            continue

        # bullet list
        if stripped.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            add_runs(p, stripped[2:])
            i += 1
            continue

        # numbered list
        mnum = re.match(r"^(\d+)\.\s+(.*)", stripped)
        if mnum:
            p = doc.add_paragraph(style="List Number")
            add_runs(p, mnum.group(2))
            i += 1
            continue

        # plain paragraph
        p = doc.add_paragraph()
        add_runs(p, stripped)
        i += 1

    doc.save(OUT)
    print(f"OK: wrote {OUT}")


if __name__ == "__main__":
    main()
